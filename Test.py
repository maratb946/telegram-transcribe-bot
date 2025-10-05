import os
import tempfile
import asyncio
from datetime import datetime
import language_tool_python
from aiogram import Bot, Dispatcher, types
from aiogram.types import Message, CallbackQuery, FSInputFile
from aiogram.filters import Command, StateFilter
from aiogram.fsm.state import State, StatesGroup
from aiogram.fsm.context import FSMContext
from aiogram.fsm.storage.memory import MemoryStorage
from aiogram.utils.keyboard import InlineKeyboardBuilder
from faster_whisper import WhisperModel

# Для DOCX
from docx import Document

# Для PDF
import pdfkit
import platform

# === Настройки ===
import os
BOT_TOKEN = os.getenv("BOT_TOKEN")
bot = Bot(token=BOT_TOKEN)
storage = MemoryStorage()
dp = Dispatcher(storage=storage)

# Путь к wkhtmltopdf (если не в PATH)
WKHTMLTOPDF_PATH = None  # Например: r"C:\Program Files\wkhtmltopdf\bin\wkhtmltopdf.exe"

# Загружаем модели
model = WhisperModel("base", device="cpu", compute_type="int8")
tool = language_tool_python.LanguageTool('ru')

# === Состояния ===
class TranscribeState(StatesGroup):
    waiting_for_correction = State()
    waiting_for_format = State()
    raw_text = State()
    detected_lang = State()
    temp_ogg = State()
    chat_id = State()
    progress_msg_id = State()
    final_text = State()

# === Вспомогательные функции ===
def correct_text(text: str, lang: str = 'ru') -> str:
    try:
        if lang != tool.language:
            tool.language = lang
        matches = tool.check(text)
        return language_tool_python.utils.correct(text, matches)
    except Exception as e:
        print(f"Ошибка коррекции: {e}")
        return text

def text_to_docx(text: str) -> str:
    doc = Document()
    doc.add_heading('Расшифровка аудио', 0)
    doc.add_paragraph(text)
    doc.add_paragraph(f"\n— Сгенерировано: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
        doc.save(f.name)
        return f.name

def text_to_pdf(text: str) -> str:
    # Экранируем HTML-символы
    safe_text = text.replace("&", "&amp;").replace("<", "<").replace(">", ">")
    html = f"""
    <html>
    <head><meta charset="UTF-8"></head>
    <body>
        <h1>Расшифровка аудио</h1>
        <pre style="white-space: pre-wrap; font-family: Arial, sans-serif;">{safe_text}</pre>
        <p><i>— Сгенерировано: {datetime.now().strftime('%Y-%m-%d %H:%M')}</i></p>
    </body>
    </html>
    """
    with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as f:
        options = {
            'page-size': 'A4',
            'margin-top': '0.75in',
            'margin-right': '0.75in',
            'margin-bottom': '0.75in',
            'margin-left': '0.75in',
            'encoding': "UTF-8",
        }
        config = pdfkit.configuration(wkhtmltopdf=WKHTMLTOPDF_PATH) if WKHTMLTOPDF_PATH else None
        pdfkit.from_string(html, f.name, options=options, configuration=config)
        return f.name

# === Обработчики ===
@dp.message(Command("start"))
async def start_handler(message: Message):
    await message.answer("Привет! Отправь голосовое или аудиофайл — я расшифрую его.")

@dp.message(StateFilter(None))
async def handle_audio(message: Message, state: FSMContext):
    if not (message.voice or message.audio):
        await message.answer("Пожалуйста, отправь голосовое сообщение или аудиофайл.")
        return

    progress_msg = await message.answer("📥 Получаю аудио...")

    file = await bot.get_file((message.voice or message.audio).file_id)
    file_path = file.file_path

    with tempfile.NamedTemporaryFile(suffix=".ogg", delete=False) as temp_audio:
        temp_ogg = temp_audio.name

    try:
        await bot.download_file(file_path, temp_ogg)
        await progress_msg.edit_text("🔄 Распознаю речь...")

        segments, info = model.transcribe(temp_ogg, beam_size=5)
        raw_text = " ".join([segment.text for segment in segments]).strip()

        if not raw_text:
            await progress_msg.edit_text("❌ Не удалось распознать речь.")
            os.unlink(temp_ogg)
            return

        await state.update_data(
            raw_text=raw_text,
            detected_lang=info.language,
            temp_ogg=temp_ogg,
            chat_id=message.chat.id,
            progress_msg_id=progress_msg.message_id
        )

        builder = InlineKeyboardBuilder()
        builder.button(text="✅ Да", callback_data="corr_yes")
        builder.button(text="❌ Нет", callback_data="corr_no")
        builder.adjust(2)

        await progress_msg.edit_text(
            f"Расшифровка готова! Язык: {info.language.upper()}\n\nХочешь, чтобы я исправил ошибки?",
            reply_markup=builder.as_markup()
        )
        await state.set_state(TranscribeState.waiting_for_correction)

    except Exception as e:
        await progress_msg.edit_text(f"⚠️ Ошибка: {str(e)}")
        if os.path.exists(temp_ogg):
            os.unlink(temp_ogg)

# === Выбор коррекции ===
@dp.callback_query(TranscribeState.waiting_for_correction, lambda c: c.data in ["corr_yes", "corr_no"])
async def process_correction_choice(callback: CallbackQuery, state: FSMContext):
    await callback.answer()
    data = await state.get_data()
    raw_text = data["raw_text"]
    detected_lang = data["detected_lang"]

    final_text = correct_text(raw_text, lang=detected_lang) if callback.data == "corr_yes" else raw_text
    await state.update_data(final_text=final_text)

    # Выбор формата
    builder = InlineKeyboardBuilder()
    builder.button(text="📩 Сообщение", callback_data="fmt_msg")
    builder.button(text="📄 TXT", callback_data="fmt_txt")
    builder.button(text="📝 DOCX", callback_data="fmt_docx")
    builder.button(text="📄 PDF", callback_data="fmt_pdf")
    builder.adjust(2)

    await callback.message.edit_text("В каком формате прислать расшифровку?", reply_markup=builder.as_markup())
    await state.set_state(TranscribeState.waiting_for_format)

# === Выбор формата ===
@dp.callback_query(TranscribeState.waiting_for_format)
async def process_format_choice(callback: CallbackQuery, state: FSMContext):
    await callback.answer()
    data = await state.get_data()
    final_text = data["final_text"]
    chat_id = data["chat_id"]
    temp_ogg = data["temp_ogg"]
    progress_msg_id = data["progress_msg_id"]

    try:
        await bot.delete_message(chat_id=chat_id, message_id=progress_msg_id)
    except:
        pass

    try:
        if callback.data == "fmt_msg":
            if len(final_text) <= 4096:
                await bot.send_message(chat_id, f"✅ Расшифровка:\n\n{final_text}")
            else:
                parts = [final_text[i:i+4096] for i in range(0, len(final_text), 4096)]
                await bot.send_message(chat_id, "✅ Расшифровка (разделена на части):")
                for part in parts:
                    await bot.send_message(chat_id, part)

        elif callback.data == "fmt_txt":
            with tempfile.NamedTemporaryFile(mode="w", suffix=".txt", delete=False, encoding="utf-8") as f:
                f.write(final_text)
                path = f.name
            await bot.send_document(chat_id, FSInputFile(path, filename="расшифровка.txt"))
            os.unlink(path)

        elif callback.data == "fmt_docx":
            path = text_to_docx(final_text)
            await bot.send_document(chat_id, FSInputFile(path, filename="расшифровка.docx"))
            os.unlink(path)

        elif callback.data == "fmt_pdf":
            path = text_to_pdf(final_text)
            await bot.send_document(chat_id, FSInputFile(path, filename="расшифровка.pdf"))
            os.unlink(path)

    except Exception as e:
        await bot.send_message(chat_id, f"❌ Ошибка при создании файла: {str(e)}")

    # Очистка
    if os.path.exists(temp_ogg):
        os.unlink(temp_ogg)
    await state.clear()

# === Запуск ===
async def main():
    print("Бот запущен с выбором коррекции и формата (TXT/DOCX/PDF/Сообщение)!")
    await dp.start_polling(bot)

if __name__ == "__main__":
    asyncio.run(main())